# # app.py
# import io
# import json
# import tempfile
# import os
# from datetime import datetime
# from typing import List

# import streamlit as st
# import fitz       # PyMuPDF
# import docx       # python-docx
# from langchain_core.documents import Document

# from rag_testcases.config import (
#     QUERIES,
#     validate_config,
#     DEBUG_MODE,
# )
# from rag_testcases.models import load_vectorstore, load_reranker
# from rag_testcases.retrieval import build_bm25_corpus
# from rag_testcases.pipeline import run_single_query
# from rag_testcases.ingestion import build_or_update_faiss

# # Import the actual function from extractor
# from extractor.qwen_vl_extractor import extract_from_image


# def extract_text_from_image_bytes(image_bytes: bytes) -> str:
#     """
#     Local adapter for UI layer.
#     Saves bytes to a temporary file and calls the actual extract_from_image function.
#     """
#     # Create a temporary file
#     with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
#         tmp_file.write(image_bytes)
#         tmp_path = tmp_file.name
    
#     try:
#         # Use the actual function which expects a file path
#         result = extract_from_image(tmp_path)
#         return result
#     finally:
#         # Clean up the temporary file
#         try:
#             os.unlink(tmp_path)
#         except:
#             pass


# # ---------- In-memory converters for uploaded files ----------

# def uploaded_pdf_to_docs(file_bytes: bytes, name: str) -> List[Document]:
#     docs: List[Document] = []
#     with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
#         for page_num, page in enumerate(pdf, start=1):
#             text = page.get_text("text")
#             if text and text.strip():
#                 docs.append(
#                     Document(
#                         page_content=text,
#                         metadata={
#                             "source": name,
#                             "page": page_num,
#                             "modality": "pdf_text",
#                         },
#                     )
#                 )
#             else:
#                 # Likely scanned/image-only page → use Qwen-VL
#                 pix = page.get_pixmap()
#                 img_bytes = pix.tobytes("png")
#                 img_text = extract_text_from_image_bytes(img_bytes)
#                 if img_text and img_text.strip():
#                     docs.append(
#                         Document(
#                             page_content=img_text,
#                             metadata={
#                                 "source": name,
#                                 "page": page_num,
#                                 "modality": "pdf_image",
#                             },
#                         )
#                     )
#     return docs


# def uploaded_docx_to_docs(file_bytes: bytes, name: str) -> List[Document]:
#     docs: List[Document] = []
#     bio = io.BytesIO(file_bytes)
#     d = docx.Document(bio)
#     paras = [p.text for p in d.paragraphs if p.text.strip()]
#     if paras:
#         docs.append(
#             Document(
#                 page_content="\n".join(paras),
#                 metadata={
#                     "source": name,
#                     "modality": "docx",
#                 },
#             )
#         )
#     return docs


# def uploaded_image_to_docs(file_bytes: bytes, name: str) -> List[Document]:
#     text = extract_text_from_image_bytes(file_bytes)
#     if not text or not text.strip():
#         return []
#     return [
#         Document(
#             page_content=text,
#             metadata={
#                 "source": name,
#                 "modality": "image",
#             },
#         )
#     ]


# # ---------- Cached pipeline init ----------

# @st.cache_resource(show_spinner=True)
# def init_pipeline():
#     """Load config, vectorstore, BM25, and reranker once per process."""
#     validate_config()
#     vs = load_vectorstore()
#     bm25_model, bm25_docs = build_bm25_corpus(vs)
#     reranker = load_reranker()
#     return vs, bm25_model, bm25_docs, reranker


# # ---------- Streamlit App ----------

# def main():
#     st.set_page_config(
#         page_title="RAG Test Case Generator",
#         layout="wide",
#     )

#     st.title("RAG Test Case Generator")
#     st.markdown(
#         "Generate **use cases & test cases** from your indexed PRD/screenshots "
#         "using a hybrid RAG pipeline."
#     )

#     # Sidebar controls
#     with st.sidebar:
#         st.header("Run settings")

#         uploaded_files = st.file_uploader(
#             "Upload files to extend the FAISS index (optional)",
#             type=["pdf", "docx", "png", "jpg", "jpeg"],
#             accept_multiple_files=True,
#             help="Uploaded files are embedded and added on top of the existing index.",
#         )

#         mode = st.radio(
#             "Query source",
#             options=["Predefined", "Custom"],
#             index=0,
#             help="Use one of the built‑in demo queries or type your own.",
#         )

#         if mode == "Predefined":
#             query = st.selectbox(
#                 "Select predefined query",
#                 options=QUERIES,
#                 index=0,
#             )
#         else:
#             query = st.text_area(
#                 "Custom query",
#                 value="Generate positive, negative, and boundary test cases for ...",
#                 height=140,
#             )

#         run_button = st.button("Run generator", type="primary")

#     # Wait for click
#     if not run_button:
#         st.info("Upload files if needed, select/enter a query, then click **Run generator**.")
#         return

#     if not query or not query.strip():
#         st.error("Query cannot be empty.")
#         return

#     # Init pipeline
#     with st.spinner("Loading vectorstore and reranker..."):
#         vs, bm25_model, bm25_docs, reranker = init_pipeline()

#     # If any files were uploaded, extend FAISS index on disk + in memory
#     if uploaded_files:
#         new_docs: List[Document] = []
#         for f in uploaded_files:
#             content = f.read()
#             lower_name = f.name.lower()

#             if lower_name.endswith(".pdf"):
#                 new_docs.extend(uploaded_pdf_to_docs(content, f.name))
#             elif lower_name.endswith(".docx"):
#                 new_docs.extend(uploaded_docx_to_docs(content, f.name))
#             elif lower_name.endswith((".png", ".jpg", ".jpeg")):
#                 new_docs.extend(uploaded_image_to_docs(content, f.name))

#         if new_docs:
#             with st.spinner(f"Indexing {len(new_docs)} new chunks into FAISS..."):
#                 vs = build_or_update_faiss(new_docs, existing_vs=vs)
#                 bm25_model, bm25_docs = build_bm25_corpus(vs)

#             st.success(f"Added {len(new_docs)} chunks from uploaded files to the index.")
#         else:
#             st.warning("No text could be extracted from uploaded files.")

#     # Show query
#     st.write("### Query")
#     st.code(query)

#     # Run pipeline
#     with st.spinner("Running RAG pipeline (retrieve → rerank → generate JSON)..."):
#         result = run_single_query(vs, bm25_model, bm25_docs, reranker, query)

#     # Show result
#     st.write("### Generated JSON")
#     st.json(result)

#     # Download JSON
#     ts = datetime.now().strftime("%Y%m%d_%H%M%S")
#     filename = f"testcases_{ts}.json"
#     file_bytes = json.dumps(result, indent=2, ensure_ascii=False).encode("utf-8")

#     st.download_button(
#         label=f"Download JSON ({filename})",
#         data=file_bytes,
#         file_name=filename,
#         mime="application/json",
#     )

#     if DEBUG_MODE:
#         st.caption(
#             "DEBUG_MODE is ON; see server logs for detailed retrieval and generation traces."
#         )


# if __name__ == "__main__":
#     main()

# app.py
import io
import json
import tempfile
import os
from datetime import datetime
from typing import List

import streamlit as st
import fitz       # PyMuPDF
import docx       # python-docx
from langchain_core.documents import Document

from rag_testcases.config import (
    QUERIES,
    validate_config,
    DEBUG_MODE,
)
from rag_testcases.models import load_vectorstore, load_reranker
from rag_testcases.retrieval import build_bm25_corpus
from rag_testcases.pipeline import run_single_query
from rag_testcases.ingestion import build_or_update_faiss

# Import the actual function from extractor
from extractor.qwen_vl_extractor import extract_from_image


def extract_text_from_image_bytes(image_bytes: bytes) -> str:
    """
    Local adapter for UI layer.
    Saves bytes to a temporary file and calls the actual extract_from_image function.
    """
    # Create a temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        tmp_file.write(image_bytes)
        tmp_path = tmp_file.name
    
    try:
        # Use the actual function which expects a file path
        result = extract_from_image(tmp_path)
        return result
    finally:
        # Clean up the temporary file
        try:
            os.unlink(tmp_path)
        except:
            pass


# ---------- In-memory converters for uploaded files ----------

def uploaded_pdf_to_docs(file_bytes: bytes, name: str) -> List[Document]:
    docs: List[Document] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text")
            if text and text.strip():
                docs.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": name,
                            "page": page_num,
                            "modality": "pdf_text",
                        },
                    )
                )
            else:
                # Likely scanned/image-only page → use Qwen-VL
                pix = page.get_pixmap()
                img_bytes = pix.tobytes("png")
                img_text = extract_text_from_image_bytes(img_bytes)
                if img_text and img_text.strip():
                    docs.append(
                        Document(
                            page_content=img_text,
                            metadata={
                                "source": name,
                                "page": page_num,
                                "modality": "pdf_image",
                            },
                        )
                    )
    return docs


def uploaded_docx_to_docs(file_bytes: bytes, name: str) -> List[Document]:
    docs: List[Document] = []
    bio = io.BytesIO(file_bytes)
    d = docx.Document(bio)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    if paras:
        docs.append(
            Document(
                page_content="\n".join(paras),
                metadata={
                    "source": name,
                    "modality": "docx",
                },
            )
        )
    return docs


def uploaded_image_to_docs(file_bytes: bytes, name: str) -> List[Document]:
    text = extract_text_from_image_bytes(file_bytes)
    if not text or not text.strip():
        return []
    return [
        Document(
            page_content=text,
            metadata={
                "source": name,
                "modality": "image",
            },
        )
    ]


# ---------- Cached pipeline init ----------

@st.cache_resource(show_spinner=True)
def init_pipeline():
    """Load config, vectorstore, BM25, and reranker once per process."""
    validate_config()
    vs = load_vectorstore()
    bm25_model, bm25_docs = build_bm25_corpus(vs)
    reranker = load_reranker()
    return vs, bm25_model, bm25_docs, reranker


# ---------- Streamlit App ----------

def main():
    st.set_page_config(
        page_title="RAG Test Case Generator",
        layout="wide",
    )

    st.title("RAG Test Case Generator")
    st.markdown(
        "Generate **use cases & test cases** from your indexed PRD/screenshots "
        "using a hybrid RAG pipeline."
    )

    # Sidebar controls
    with st.sidebar:
        st.header("Run settings")

        uploaded_files = st.file_uploader(
            "Upload files to extend the FAISS index (optional)",
            type=["pdf", "docx", "png", "jpg", "jpeg"],
            accept_multiple_files=True,
            help="Uploaded files are embedded and added on top of the existing index.",
        )

        mode = st.radio(
            "Query source",
            options=["Predefined", "Custom"],
            index=0,
            help="Use one of the built‑in demo queries or type your own.",
        )

        if mode == "Predefined":
            query = st.selectbox(
                "Select predefined query",
                options=QUERIES,
                index=0,
            )
        else:
            query = st.text_area(
                "Custom query",
                value="Generate positive, negative, and boundary test cases for ...",
                height=140,
            )

        run_button = st.button("Run generator", type="primary")

    # Wait for click
    if not run_button:
        st.info("Upload files if needed, select/enter a query, then click **Run generator**.")
        return

    if not query or not query.strip():
        st.error("Query cannot be empty.")
        return

    # Init pipeline
    with st.spinner("Loading vectorstore and reranker..."):
        vs, bm25_model, bm25_docs, reranker = init_pipeline()

    # Create columns for better layout
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # If any files were uploaded, extend FAISS index on disk + in memory
        if uploaded_files:
            st.write("### File Upload Status")
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            new_docs: List[Document] = []
            total_files = len(uploaded_files)
            
            for i, f in enumerate(uploaded_files):
                # Update progress
                progress = (i / total_files) * 100
                progress_bar.progress(progress / 100)
                status_text.text(f"Processing {i+1}/{total_files}: {f.name}")
                
                content = f.read()
                lower_name = f.name.lower()

                if lower_name.endswith(".pdf"):
                    file_docs = uploaded_pdf_to_docs(content, f.name)
                    new_docs.extend(file_docs)
                    st.info(f"📄 {f.name}: Extracted {len(file_docs)} pages")
                elif lower_name.endswith(".docx"):
                    file_docs = uploaded_docx_to_docs(content, f.name)
                    new_docs.extend(file_docs)
                    st.info(f"📝 {f.name}: Extracted {len(file_docs)} document(s)")
                elif lower_name.endswith((".png", ".jpg", ".jpeg")):
                    file_docs = uploaded_image_to_docs(content, f.name)
                    new_docs.extend(file_docs)
                    st.info(f"🖼️ {f.name}: Extracted {len(file_docs)} image(s)")
                    
                # Reset file pointer for next iteration if needed
                f.seek(0)
            
            # Final progress update
            progress_bar.progress(1.0)
            status_text.text(f"Processed all {total_files} files")
            
            if new_docs:
                st.write(f"### Embedding Status")
                embedding_status = st.empty()
                embedding_status.info(f"Embedding {len(new_docs)} new chunks into FAISS...")
                
                # Show a spinner during embedding
                with st.spinner(f"Creating embeddings for {len(new_docs)} chunks..."):
                    vs = build_or_update_faiss(new_docs, existing_vs=vs)
                    bm25_model, bm25_docs = build_bm25_corpus(vs)
                
                embedding_status.success(f"✅ Successfully embedded {len(new_docs)} chunks")
                st.balloons()  # Celebrate success
                
                # Show summary
                with st.expander("📊 Upload Summary", expanded=True):
                    col_a, col_b = st.columns(2)
                    with col_a:
                        st.metric("Files Uploaded", total_files)
                    with col_b:
                        st.metric("Chunks Added", len(new_docs))
            else:
                st.warning("⚠️ No text could be extracted from uploaded files.")
            
            # Clear progress indicators
            progress_bar.empty()
            status_text.empty()
            
        else:
            st.info("ℹ️ No new files uploaded. Using existing index.")
    
    with col2:
        st.write("### Current Index Status")
        try:
            # Try to get some stats from the vectorstore
            st.metric("Vector Store", "FAISS")
            # Note: FAISS doesn't easily expose count, but we can estimate
            st.metric("BM25 Docs", f"{len(bm25_docs) if bm25_docs else 'N/A'}")
        except:
            pass
    
    # Divider
    st.divider()
    
    # Show query in a prominent box
    st.write("## 📝 Query")
    st.info(f"**Query:** {query}")
    
    # Run pipeline with better feedback
    st.write("## ⚙️ Pipeline Execution")
    
    with st.expander("View Pipeline Steps", expanded=True):
        step1, step2, step3, step4 = st.columns(4)
        
        with step1:
            st.markdown("### 1️⃣")
            st.markdown("**Retrieval**")
        with step2:
            st.markdown("### 2️⃣")
            st.markdown("**Reranking**")
        with step3:
            st.markdown("### 3️⃣")
            st.markdown("**Generation**")
        with step4:
            st.markdown("### 4️⃣")
            st.markdown("**Output**")
    
    # Run the actual pipeline
    pipeline_status = st.empty()
    pipeline_status.info("Starting RAG pipeline...")
    
    # Use a container for the results
    results_container = st.container()
    
    with st.spinner("Running RAG pipeline (retrieve → rerank → generate JSON)..."):
        pipeline_status.info("🔍 Retrieving relevant documents...")
        result = run_single_query(vs, bm25_model, bm25_docs, reranker, query)
        pipeline_status.success("✅ Pipeline completed successfully!")
    
    # Display results
    with results_container:
        st.write("## 📋 Generated Results")
        
        # Show result in JSON format
        tab1, tab2 = st.tabs(["📄 JSON Output", "📊 Summary"])
        
        with tab1:
            st.write("### Generated JSON")
            st.json(result)
            
            # Download JSON button
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"testcases_{ts}.json"
            file_bytes = json.dumps(result, indent=2, ensure_ascii=False).encode("utf-8")
            
            st.download_button(
                label=f"📥 Download JSON ({filename})",
                data=file_bytes,
                file_name=filename,
                mime="application/json",
                icon="📥",
                use_container_width=True
            )
        
        with tab2:
            # Try to extract some summary info from the result
            try:
                if isinstance(result, dict):
                    total_cases = 0
                    if "positive" in result:
                        total_cases += len(result.get("positive", []))
                    if "negative" in result:
                        total_cases += len(result.get("negative", []))
                    if "boundary" in result:
                        total_cases += len(result.get("boundary", []))
                    
                    col_a, col_b, col_c = st.columns(3)
                    with col_a:
                        st.metric("Positive Cases", len(result.get("positive", [])))
                    with col_b:
                        st.metric("Negative Cases", len(result.get("negative", [])))
                    with col_c:
                        st.metric("Boundary Cases", len(result.get("boundary", [])))
                    
                    st.metric("Total Test Cases", total_cases)
                else:
                    st.info("Result is not in expected dictionary format.")
            except:
                st.info("Could not extract summary metrics from result.")
    
    # Debug mode notice
    if DEBUG_MODE:
        st.divider()
        st.warning(
            "**DEBUG_MODE is ON** - see server logs for detailed retrieval and generation traces."
        )


if __name__ == "__main__":
    main()