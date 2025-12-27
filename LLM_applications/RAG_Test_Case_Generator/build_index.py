# build_index.py
#
# Offline ingestion: read PDFs, DOCX, and Booking screenshots,
# extract text (incl. via Qwen-VL), chunk, and build a local FAISS index.

import os
import uuid
import time

from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from docx import Document as DocxDocument
import fitz  # PyMuPDF

from extractor.qwen_vl_extractor import (
    extract_image_chunks,
    clean_text,
    chunk_text,
)

# ==========
# CONSTANTS
# ==========
DATA_DIR = "./sample_data"
FAISS_DIR = "./faiss_store"
EMBED_MODEL = "all-mpnet-base-v2"


def extract_text_from_pdf(path: str):
    doc = fitz.open(path)
    pages = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            pages.append((page.number + 1, text))
    return pages


def extract_text_from_docx(path: str):
    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return text.strip()


def iter_files(data_dir: str):
    for root, _, files in os.walk(data_dir):
        for f in files:
            yield os.path.join(root, f)


def build_index():
    documents = []
    start_time = time.time()

    files = list(iter_files(DATA_DIR))
    print(f"📂 Found {len(files)} files in {DATA_DIR}")

    for path in files:
        ext = path.lower().split(".")[-1]
        file_name = os.path.basename(path)

        if ext in ["png", "jpg", "jpeg"]:
            # image → Qwen-VL → chunks
            chunks, modality = extract_image_chunks(path)
            file_type = "image"
            chunk_items = [(None, c) for c in chunks]

        elif ext == "pdf":
            pages = extract_text_from_pdf(path)
            modality = "text"
            file_type = "pdf"
            chunk_items = []
            for page_no, text in pages:
                cleaned = clean_text(text)
                for c in chunk_text(cleaned):
                    chunk_items.append((page_no, c))

        elif ext == "docx":
            text = extract_text_from_docx(path)
            if not text:
                print(f"⚠️ Skipped empty DOCX {file_name}")
                continue
            modality = "text"
            file_type = "docx"
            cleaned = clean_text(text)
            chunk_items = [(None, c) for c in chunk_text(cleaned)]

        else:
            print(f"⚠️ Unsupported file type: {file_name}")
            continue

        if not chunk_items:
            print(f"⚠️ No chunks for {file_name}")
            continue

        print(f"✅ {file_name} → {len(chunk_items)} chunks")
        for page_no, chunk_text_content in chunk_items:
            documents.append(
                Document(
                    page_content=chunk_text_content,
                    metadata={
                        "source": file_name,
                        "file_path": path,
                        "modality": modality,
                        "file_type": file_type,
                        "page": page_no,
                        "chunk_id": str(uuid.uuid4()),
                    },
                )
            )

    if not documents:
        raise RuntimeError("❌ No valid documents to index")

    print(f"📄 Total chunks: {len(documents)}")

    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
    t0 = time.time()
    vectorstore = FAISS.from_documents(documents, embeddings)
    build_latency = time.time() - t0
    print(f"⏱️ Vector build time: {build_latency:.2f}s")

    os.makedirs(FAISS_DIR, exist_ok=True)
    vectorstore.save_local(FAISS_DIR)

    print(f"🎯 FAISS index saved to '{FAISS_DIR}' in {time.time() - start_time:.2f}s")


if __name__ == "__main__":
    build_index()
