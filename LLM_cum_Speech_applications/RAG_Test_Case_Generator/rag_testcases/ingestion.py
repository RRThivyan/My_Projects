# rag_testcases/ingestion.py
import os
from typing import List, Optional
import tempfile

from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
import fitz  # PyMuPDF
import docx  # python-docx

from .config import FAISS_DIR, EMBED_MODEL

# Import the actual function from extractor
from extractor.qwen_vl_extractor import extract_from_image


def extract_text_from_image(image_bytes: bytes) -> str:
    """
    Local adapter for ingestion/app.
    Saves bytes to a temporary file and calls the actual extract_from_image function.
    """
    # Create a temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        tmp_file.write(image_bytes)
        tmp_path = tmp_file.name
    
    try:
        # Use the actual function which expects a file path
        result = extract_from_image(tmp_path)
        return result
    finally:
        # Clean up the temporary file
        try:
            os.unlink(tmp_path)
        except:
            pass


def pdf_to_docs(path: str) -> List[Document]:
    """Extract text from a PDF; use Qwen-VL for image-only pages."""
    docs: List[Document] = []
    with fitz.open(path) as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text")
            if text and text.strip():
                docs.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": os.path.basename(path),
                            "page": page_num,
                            "modality": "pdf_text",
                        },
                    )
                )
            else:
                pix = page.get_pixmap()
                img_bytes = pix.tobytes("png")
                img_text = extract_text_from_image(img_bytes)
                if img_text and img_text.strip():
                    docs.append(
                        Document(
                            page_content=img_text,
                            metadata={
                                "source": os.path.basename(path),
                                "page": page_num,
                                "modality": "pdf_image",
                            },
                        )
                    )
    return docs


def docx_to_docs(path: str) -> List[Document]:
    docs: List[Document] = []
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    if paras:
        docs.append(
            Document(
                page_content="\n".join(paras),
                metadata={
                    "source": os.path.basename(path),
                    "modality": "docx",
                },
            )
        )
    return docs


def image_to_docs(path: str) -> List[Document]:
    """Extract text from image files using Qwen-VL."""
    # Since extract_from_image expects a file path, we can call it directly
    text = extract_from_image(path)
    if not text or not text.strip():
        return []
    return [
        Document(
            page_content=text,
            metadata={
                "source": os.path.basename(path),
                "modality": "image",
            },
        )
    ]


def load_files_to_docs(paths: List[str]) -> List[Document]:
    all_docs: List[Document] = []
    for path in paths:
        lower = path.lower()
        if lower.endswith(".pdf"):
            all_docs.extend(pdf_to_docs(path))
        elif lower.endswith(".docx"):
            all_docs.extend(docx_to_docs(path))
        elif lower.endswith((".png", ".jpg", ".jpeg")):
            all_docs.extend(image_to_docs(path))
    return all_docs


def get_embeddings():
    return HuggingFaceEmbeddings(model_name=EMBED_MODEL)


def build_or_update_faiss(
    docs: List[Document],
    existing_vs: Optional[FAISS] = None,
) -> FAISS:
    embeddings = get_embeddings()
    if existing_vs is None:
        vs = FAISS.from_documents(docs, embeddings)
    else:
        existing_vs.add_documents(docs)
        vs = existing_vs

    os.makedirs(FAISS_DIR, exist_ok=True)
    vs.save_local(FAISS_DIR)
    return vs